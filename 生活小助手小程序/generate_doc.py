# -*- coding: utf-8 -*-
"""
生活小助手小程序 - 功能清单 Word 文档生成器
- 按 5 个模块分页，每页一个主题色
- 保留 CSV 全部 7 列
- 「功能更新迭代记录」列填入源码核验结果
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============ 工具函数 ============
def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run_font(run, size=10.5, color_hex='000000', bold=False, font_name='微软雅黑'):
    """设置 run 字体（含中文 eastAsia）"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:eastAsia'), font_name)
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)


def set_cell_borders(cell, color_hex='BFBFBF', size='4'):
    """设置单元格四边边框"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement('w:' + edge)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def write_cell(cell, text, size=10.5, color_hex='000000', bold=False,
               align='left', bg=None, font_name='微软雅黑'):
    """向单元格写入文本（支持多行：用 \n 分隔）"""
    cell.text = ''
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        set_cell_bg(cell, bg)
    set_cell_borders(cell)
    lines = str(text).split('\n') if text else ['']
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }[align]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        set_run_font(run, size=size, color_hex=color_hex, bold=bold, font_name=font_name)


def set_col_widths(table, widths_cm):
    """强制设置每列宽度（cm 列表）"""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            row.cells[idx].width = Cm(w)


# ============ 数据 ============
HEADERS = ['序号', '所属页面/模块', '功能点名称', '功能详细描述',
           '待优化/待完善问题', '功能更新迭代记录', '备注']
COL_WIDTHS = [1.1, 1.9, 2.6, 5.6, 5.4, 5.0, 1.4]  # 合计 23cm（横向 A4 可用宽度约 24cm）

MODULES = [
    {
        "name": "主页面（首页 index）",
        "page_label": "主页面",
        "primary": "2E86AB",   # 深蓝
        "light": "EAF3F8",     # 浅蓝
        "warn": "FDF2D8",      # 浅黄（待优化行高亮）
        "items": [
            (1, "账户板块",
             "支持用户上传头像、自主编辑个人基础信息，展示用户基础身份标识",
             "审核待优化：当前强制登录才能浏览，需调整为用户可先浏览功能、触发写操作时再引导登录",
             "【已优化】源码 getHomeData() 已实现游客浏览（guestHome 兜底），requireLogin() 仅在写操作时弹窗引导登录，浏览态无需强制登录"),
            (2, "设置入口-页面管理",
             "可对主页面底部导航栏进行管理，隐藏不需要的功能入口，自定义首页展示模块",
             "无基础功能缺陷，可按需扩展隐藏/显示的模块范围",
             "【已实现】settings.js 实现 PAGE_INFO 配置 + toggleHide 隐藏 + moveUp/moveDown 排序 + tabOrder 持久化，5 个 tab 可独立显隐与排序"),
            (3, "设置入口-色调调整",
             "支持对页面局部色调进行微调，修改配色风格",
             "待优化：需实现色调整体统一调整，用户修改后全页面配色同步生效，而非仅局部调整",
             "【已优化】源码 theme.apply(name) 已实现全局主题同步：导航栏 + 所有活跃页面 + TabBar 实时刷新；theme.getThemeList() 提供多套主题切换"),
            (4, "设置入口-意见反馈",
             "用户可提交意见反馈，内容可上传至云端后台，供开发者查看",
             "无基础功能缺陷，可后续迭代反馈分类、提交状态提示等功能",
             "【已实现】settings.js saveFeedback() 提交至云端 feedback 集合（含 userId/account/createdAt），云端不可用时本地兜底保存"),
            (5, "首页自定义板块",
             "仅固定展示预设的功能板块，无自定义调整能力",
             "待优化：支持用户自主添加记录内容/方向的板块，可自定义功能偏向、编辑卡片标题，实现首页模块化定制",
             "【已优化】源码已完整实现：openManage 板块管理（拖拽排序/重命名/隐藏/删除）+ openCreate 新建板块（支持 todo/checkin/task/number/custom 5 种类型）"),
            (6, "今日待办",
             "展示用户待办任务列表，支持勾选标记任务完成状态",
             "无基础功能缺陷，可后续扩展任务分类、截止日期提醒功能",
             "【已实现】index.js todo 模块完整：addTodo/toggleTodo/deleteTodo + 进度条 + 完成率统计；板块内任务管理闭环"),
            (7, "每日打卡",
             "提供每日打卡功能，记录用户每日打卡记录，形成连续打卡记录",
             "无基础功能缺陷，可后续扩展打卡补签、打卡统计看板功能",
             "【已实现】index.js checkin 模块完整：30 天打卡网格 + 连续天数 streak + 目标天数 slider + 月度进度；点击任意日期即可打卡/取消"),
            (8, "近期任务",
             "支持设置任务重要程度、截止日期，展示待办任务的时间与优先级",
             "待优化：任务到达截止日期时，需在微信内推送消息提醒用户",
             "【已优化】源码 toggleTaskRemind() 已实现订阅消息提醒：取 openid → 请求订阅授权 → 写入提醒标记；由云函数 taskReminder 定时触发当天推送"),
            (9, "体重记录",
             "记录用户体重数据，每条记录同步显示记录日期，形成体重变化记录",
             "功能异常：当前体重记录无法正常显示，需修复数据渲染问题",
             "【源码完整·待复测】computeModules() 中 weight 分支与 refreshModal() weight 分支渲染逻辑完整（含趋势箭头/环比 diff），源码层面无缺陷，疑为运行时数据或环境问题，建议真机复测定位"),
            (10, "首页自定义板块-数值记录",
             "支持自定义单位的数值记录（身高/个数/步数等），形成变化趋势",
             "无基础功能缺陷，可后续扩展图表化展示",
             "【已实现】index.js number 模块完整：支持 10 种预设单位（kg/cm/m/个/次/升/分钟/元/步/%）+ 自定义单位 + 备注 + 趋势 diff"),
        ],
    },
    {
        "name": "记账页面（account）",
        "page_label": "记账页面",
        "primary": "E76F51",   # 暖橙
        "light": "FDEEE6",     # 浅橙
        "warn": "FDF2D8",
        "items": [
            (11, "余额板块",
             "支持用户手动修改账户余额，自定义调整初始余额数值",
             "无基础功能缺陷，可后续扩展余额修改日志、操作记录追溯功能",
             "【已实现】account.js editBalance/confirmBalance 支持个人/家庭双账户余额修改，按差额折算到期初余额，避免重复计算"),
            (12, "支出统计",
             "新增支出记录后，自动同步更新账户余额，展示支出明细与统计",
             "核心缺陷：仅支持支出记录录入，无收入录入功能，余额计算逻辑闭环断裂",
             "【已优化】源码 openAddModal() 的 recordType 已支持 'income'，calculateAccountBalance() 按 type===income 正向累加，收支闭环已打通；含今日/本月支出/收入统计"),
            (13, "每日记账",
             "支持选择消费分类、选择记账时间，完成收支记录的手动录入",
             "待优化：新增语音输入、AI 智能记账/自动记账功能，支持语音口述自动生成记账记录",
             "【已实现基础】account.js 支持选分类（食/行/购等）+ 选日期 + 个人/家庭维度；【未实现】语音输入与 AI 自动记账能力尚需接入"),
            (14, "刚需必买/愿望清单",
             "记录内容不会立刻同步余额，勾选后才会触发余额变动，同时自动添加对应记账记录",
             "待优化：当前自动标记为家庭维度，无法选择个人维度，需对接标签体系实现个人/家庭维度可选切换",
             "【部分优化】wish 愿望清单已支持 ownerType（personal/family）可选；necessity 刚需必买仍固定 family 维度，需补 ownerType 字段对齐"),
            (15, "消费统计图表",
             "可按预设时间维度分类查看消费数据，生成可视化统计图表",
             "待优化：支持用户自由自定义选择起止日期，筛选对应区间的收支数据并重绘图表",
             "【已优化】account.js drawChart() 圆环图已支持 chartPeriod=today/week/month/halfyear/custom，custom 模式下 chartStart/chartEnd 自由选起止日期；支持收支切换 + 个人/家庭/全部维度"),
        ],
    },
    {
        "name": "菜谱页面（recipe）",
        "page_label": "菜谱页面",
        "primary": "2A9D8F",   # 青绿
        "light": "E4F4F1",     # 浅青
        "warn": "FDF2D8",
        "items": [
            (16, "我的食材",
             "支持用户手动添加已有的食材，形成个人食材库存列表",
             "功能审核中，需完善功能逻辑、适配平台审核要求，优化食材管理流程",
             "【已实现】recipe.js openIngredientModal/confirmAddIngredient/removeIngredient 完整：支持食材分类（vegetable 等）+ 去重 + 删除；菜谱可按现有食材反查可做菜品"),
            (17, "家常菜谱-基础管理",
             "支持查看全量菜谱、现有食材可完成的菜谱，支持上传菜品实拍图、选择所需食材、记录做菜步骤",
             "无基础功能缺陷，可后续扩展菜谱分类、食材用量统计、做法分步展示功能",
             "【已实现】recipe.js openRecipeModal/chooseImage/tempIngredients/recipeSteps 完整；renderRecipes() 含 canMake 可做判定 + missing 缺料数 + tags 食材标签；recipeFilter 支持全量/可做切换"),
            (18, "家常菜谱-删除功能",
             "支持删除已创建的菜谱，当前删除按钮尺寸过小，交互体验不畅",
             "待优化：改为点击后弹出删除/取消二次确认交互框，或放大删除按钮适配手机触控区域",
             "【未优化】源码 deleteRecipe() 直接删除，无二次确认弹窗；建议接入 wx.showModal 二次确认或放大触控区域"),
        ],
    },
    {
        "name": "进度条页面（supply）",
        "page_label": "进度条页面",
        "primary": "8E7DBE",   # 紫罗兰
        "light": "EFEAF6",     # 浅紫
        "warn": "FDF2D8",
        "items": [
            (19, "进度条分类管理",
             "用户可自定义分类，按需创建任务完成进度条、物品消耗比例条，自主控制进度数值",
             "无基础功能缺陷，可后续扩展进度条颜色自定义、分类拖拽排序、进度历史回溯功能",
             "【已实现】supply.js 完整：openAddCatModal/openEditCatModal 分类增改 + deleteCat 删除（含项目归并默认分类）+ confirmAddSupply 物品创建 + 6 色颜色选择 + openEditModal 比例修改"),
            (20, "桌面快捷卡片功能",
             "无该功能，无快捷入口能力",
             "待优化：新增桌面单个功能快捷卡片，或独立快捷功能页面，实现核心功能一键直达",
             "【未实现】源码确认无桌面快捷卡片相关逻辑；需接入微信小程序「桌面快捷方式」能力（wx.addDesktopShortcut）或独立服务消息卡片"),
        ],
    },
    {
        "name": "取件码页面（pickup）",
        "page_label": "取件码页面",
        "primary": "C9607E",   # 玫粉
        "light": "F7E6EC",     # 浅粉
        "warn": "FDF2D8",
        "items": [
            (21, "驿站分类管理",
             "支持用户自定义分类不同的快递驿站，对驿站进行分组管理",
             "无基础功能缺陷，可后续扩展驿站排序、常用驿站置顶功能",
             "【已实现】pickup.js openCatModal/openEditCat/deleteCat/confirmCat 完整：驿站增改删 + 7 种马卡龙色调区分（樱花粉/薄荷绿/鹅黄/天空蓝/薰衣草/蜜桃/抹茶）+ 同名去重校验"),
            (22, "取件码记录管理",
             "支持记录、编辑、复制、删除快递取件码，关联对应驿站分类，实现取件码统一管理",
             "无基础功能缺陷，可后续扩展取件码有效期提醒、一键复制粘贴功能",
             "【已实现】pickup.js confirmAddCodes 支持批量添加（空格/逗号/换行分隔）+ toggleCode 勾选已取 + deleteCode 删除 + 进度统计（picked/total）；【未实现】一键复制粘贴可补充 wx.setClipboardData"),
            (23, "快递信息 OCR 识别",
             "无该功能，无法自动识别快递信息",
             "待优化：增加图片 OCR 识别能力，上传快递截图后，自动识别图片中的快递站点和取件码，自动录入对应分类",
             "【未实现】源码确认无 OCR 相关逻辑；可接入微信插件 OCR 或云开发 AI 能力（如腾讯云 OCR 通用印刷体识别）"),
        ],
    },
]


# ============ 文档构建 ============
def build_document(output_path):
    doc = Document()

    # ---- 页面设置：横向 A4 + 窄边距 ----
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ---- 默认样式 ----
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for mi, mod in enumerate(MODULES):
        # 模块间分页（第一页不加 page_break）
        if mi > 0:
            doc.add_page_break()

        # ---- 模块标题 ----
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(8)
        title_run = title_p.add_run('【模块 ' + str(mi + 1) + '】' + mod['name'])
        set_run_font(title_run, size=18, color_hex=mod['primary'], bold=True)

        # 副标题：功能数概览
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub_p.paragraph_format.space_after = Pt(10)
        sub_run = sub_p.add_run('共 ' + str(len(mod['items'])) + ' 个功能点  ·  主题色 #' + mod['primary'])
        set_run_font(sub_run, size=10, color_hex='808080', bold=False)

        # ---- 表格 ----
        n_rows = 1 + len(mod['items'])
        table = doc.add_table(rows=n_rows, cols=len(HEADERS))
        table.alignment = 1  # center

        # 表头行
        for ci, h in enumerate(HEADERS):
            write_cell(table.rows[0].cells[ci], h,
                       size=11, color_hex='FFFFFF', bold=True,
                       align='center', bg=mod['primary'])

        # 数据行
        for ri, item in enumerate(mod['items'], start=1):
            seq, name, desc, issue, update = item
            note = ''
            row_cells = table.rows[ri].cells
            # 待优化问题非空 → 用浅黄高亮该单元格；功能更新记录含【未实现/未优化】→ 用浅红高亮
            issue_bg = mod['warn'] if issue and issue.strip() else None
            update_bg = None
            if update:
                if '【未实现】' in update or '【未优化】' in update:
                    update_bg = 'FBE4E4'   # 浅红
                elif '【已优化】' in update or '【已实现】' in update:
                    update_bg = 'E8F5E9'   # 浅绿

            write_cell(row_cells[0], str(seq), size=10.5, bold=True, align='center')
            write_cell(row_cells[1], mod['page_label'], size=10.5, align='center',
                       bg=mod['light'])
            write_cell(row_cells[2], name, size=10.5, bold=True, align='left')
            write_cell(row_cells[3], desc, size=10, align='left')
            write_cell(row_cells[4], issue, size=10, align='left', bg=issue_bg)
            write_cell(row_cells[5], update, size=10, align='left', bg=update_bg)
            write_cell(row_cells[6], note, size=10, align='center')

        # 强制列宽
        set_col_widths(table, COL_WIDTHS)

        # 重复表头（每页重复，便于跨页阅读）
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), 'true')
        trPr.append(tblHeader)

    doc.save(output_path)
    print('OK ->', output_path)


if __name__ == '__main__':
    out = r'C:\Users\86142\Desktop\小程序\生活小助手小程序\生活小助手小程序_功能清单.docx'
    build_document(out)
